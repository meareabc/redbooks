# -*- coding: utf-8 -*-
"""数据分析器"""

import os
from datetime import datetime
from collections import Counter

from .constants import HAS_MATPLOTLIB, HAS_WORDCLOUD, HAS_DOCX


class DataAnalyzer:
    """数据分析器"""
    
    @staticmethod
    def generate_stats(df):
        """生成统计数据"""
        stats = {
            'total_notes': len(df),
            'total_likes': df['like_count'].sum() if 'like_count' in df.columns else 0,
            'avg_likes': df['like_count'].mean() if 'like_count' in df.columns else 0,
            'max_likes': df['like_count'].max() if 'like_count' in df.columns else 0,
            'total_collects': df['collect_count'].sum() if 'collect_count' in df.columns else 0,
            'total_comments': df['comment_count'].sum() if 'comment_count' in df.columns else 0,
            'image_notes': len(df[df['note_type'] == '图文']) if 'note_type' in df.columns else 0,
            'video_notes': len(df[df['note_type'] == '视频']) if 'note_type' in df.columns else 0,
        }
        return stats
    
    @staticmethod
    def generate_charts(df, output_dir):
        """生成图表"""
        if not HAS_MATPLOTLIB:
            return []
        
        import matplotlib.pyplot as plt
        
        charts = []
        os.makedirs(output_dir, exist_ok=True)
        
        plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei']
        plt.rcParams['axes.unicode_minus'] = False
        
        try:
            # 点赞分布图
            if 'like_count' in df.columns:
                fig, ax = plt.subplots(figsize=(10, 6))
                df['like_count'].hist(bins=20, ax=ax, color='#ff6b6b', edgecolor='white')
                ax.set_title('点赞数分布', fontsize=14)
                ax.set_xlabel('点赞数')
                ax.set_ylabel('笔记数量')
                chart_path = os.path.join(output_dir, 'likes_distribution.png')
                plt.savefig(chart_path, dpi=100, bbox_inches='tight')
                plt.close()
                charts.append(chart_path)
            
            # 笔记类型饼图
            if 'note_type' in df.columns:
                fig, ax = plt.subplots(figsize=(8, 8))
                type_counts = df['note_type'].value_counts()
                ax.pie(type_counts.values, labels=type_counts.index, autopct='%1.1f%%',
                       colors=['#4ecdc4', '#ff6b6b', '#ffe66d'])
                ax.set_title('笔记类型分布', fontsize=14)
                chart_path = os.path.join(output_dir, 'type_distribution.png')
                plt.savefig(chart_path, dpi=100, bbox_inches='tight')
                plt.close()
                charts.append(chart_path)
            
            # Top10点赞笔记
            if 'like_count' in df.columns and 'title' in df.columns:
                fig, ax = plt.subplots(figsize=(12, 6))
                top10 = df.nlargest(10, 'like_count')
                titles = [t[:15] + '...' if len(t) > 15 else t for t in top10['title']]
                ax.barh(range(len(top10)), top10['like_count'], color='#667eea')
                ax.set_yticks(range(len(top10)))
                ax.set_yticklabels(titles)
                ax.set_xlabel('点赞数')
                ax.set_title('Top10 热门笔记', fontsize=14)
                ax.invert_yaxis()
                chart_path = os.path.join(output_dir, 'top10_notes.png')
                plt.savefig(chart_path, dpi=100, bbox_inches='tight')
                plt.close()
                charts.append(chart_path)
                
        except Exception:
            pass
        
        return charts
    
    @staticmethod
    def generate_wordcloud(texts, output_path):
        """生成词云"""
        if not HAS_WORDCLOUD:
            return None
        
        from wordcloud import WordCloud
        import jieba
        
        try:
            all_text = ' '.join(texts)
            words = jieba.cut(all_text)
            word_list = [w for w in words if len(w) > 1]
            word_freq = Counter(word_list)
            
            wc = WordCloud(
                font_path='C:/Windows/Fonts/simhei.ttf',
                width=800,
                height=400,
                background_color='white',
                max_words=100,
                colormap='viridis'
            )
            wc.generate_from_frequencies(word_freq)
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            wc.to_file(output_path)
            return output_path
        except Exception:
            return None
    
    @staticmethod
    def generate_report(df, stats, charts, output_path, keyword):
        """生成Word分析报告"""
        if not HAS_DOCX:
            return None
        
        from docx import Document
        from docx.shared import Inches
        
        try:
            doc = Document()
            doc.add_heading(f'小红书数据分析报告 - {keyword}', 0)
            doc.add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            
            # 统计概览
            doc.add_heading('数据概览', level=1)
            table = doc.add_table(rows=4, cols=2)
            table.style = 'Table Grid'
            
            stats_items = [
                ('总笔记数', stats.get('total_notes', 0)),
                ('总点赞数', stats.get('total_likes', 0)),
                ('平均点赞', f"{stats.get('avg_likes', 0):.1f}"),
                ('最高点赞', stats.get('max_likes', 0)),
            ]
            
            for i, (label, value) in enumerate(stats_items):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
            
            # 图表
            if charts:
                doc.add_heading('数据可视化', level=1)
                for chart in charts:
                    if os.path.exists(chart):
                        doc.add_picture(chart, width=Inches(6))
                        doc.add_paragraph('')
            
            # Top10列表
            doc.add_heading('热门笔记 Top10', level=1)
            if 'like_count' in df.columns:
                top10 = df.nlargest(10, 'like_count')
                for i, row in top10.iterrows():
                    title = row.get('title', '')[:50]
                    likes = row.get('like_count', 0)
                    doc.add_paragraph(f"• {title}... (点赞 {likes})")
            
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            doc.save(output_path)
            return output_path
        except Exception:
            return None
